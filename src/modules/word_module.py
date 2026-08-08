"""
Word Document Module - OpenClaw Pattern
Handles Microsoft Word (.docx) file operations
"""
from typing import Dict, Any
import os
import re
from src.core.base_module import BaseModule

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordModule(BaseModule):
    """Module for Word document operations (create, edit, read)"""
    
    def __init__(self, name: str):
        super().__init__(name)
        self.description = "Handles Microsoft Word document operations"
        
        if not DOCX_AVAILABLE:
            print("⚠️ python-docx not installed. Install with: pip install python-docx")
            self.enabled = False
    
    async def can_handle(self, intent: Dict[str, Any]) -> bool:
        """Check if this module handles Word documents"""
        return intent.get('type') == 'word' and self.enabled
    
    async def execute(self, intent: Dict[str, Any], shared) -> str:
        """Execute Word document operation"""
        text = intent.get('text', '')
        language = intent.get('language', 'en')

        self._shared_state = shared.session_state

        text_lower = text.lower()

        # Toggle writing mode (conversation lines appended to last Word file)
        if any(word in text_lower for word in ['start writing mode', 'enable writing mode', 'modo escritura', 'activar modo escritura', 'continue writing in this file']):
            last_path = shared.session_state.get('last_word_path')
            if not last_path:
                return "No previous Word file found. Create or read one first." if language == 'en' else "No hay un archivo Word previo. Crea o lee uno primero."
            shared.session_state['stream_to_word'] = True
            return f"✍️ Writing mode enabled for: {last_path}" if language == 'en' else f"✍️ Modo escritura activado para: {last_path}"

        if any(word in text_lower for word in ['stop writing mode', 'disable writing mode', 'detener modo escritura', 'desactivar modo escritura']):
            shared.session_state['stream_to_word'] = False
            return "✍️ Writing mode disabled" if language == 'en' else "✍️ Modo escritura desactivado"

        # Fast deterministic extraction (0 tokens)
        parsed = self._extract_word_details(text, language)
        target = parsed.get('target')
        location = parsed.get('location')
        content = parsed.get('content', '')

        if not target:
            last_path = shared.session_state.get('last_word_path')
            if last_path:
                target = os.path.splitext(os.path.basename(last_path))[0]
                location = os.path.dirname(last_path)
        
        # Determine action: create, read, edit, summarize (with Spanish verb variants)
        # Spanish: crea, crees, crear, creé, escribir vs English: create, write
        if any(word in text_lower for word in ['summary', 'summarize', 'resumen', 'resumir', 'resume']):
            summary_details = self._extract_summary_details(text, language)
            source_name = summary_details.get('source') or target
            output_name = summary_details.get('output')
            summary_location = summary_details.get('location') or location

            if not source_name or not summary_location:
                command = shared.ai.understand_command(text, language)
                source_name = source_name or command.get('target')
                summary_location = summary_location or command.get('location')

            return await self._summarize_word_doc(source_name, summary_location, output_name, language, shared)

        if any(word in text_lower for word in ['create', 'crea', 'crear', 'creé', 'crees', 'escrib', 'write']):
            if not target or not location:
                command = shared.ai.understand_command(text, language)
                target = target or command.get('target')
                location = location or command.get('location')
                content = content or command.get('content', '')
            if any(word in text_lower for word in ['replace', 'overwrite', 'clear', 'remove all', 'borrar', 'reemplaz', 'elimina todo']):
                return await self._replace_word_doc(target, location, content, language)
            response = await self._create_word_doc(target, location, content, language)
            if any(phrase in text_lower for phrase in ['inside this file put this text', 'inside this file', 'in this file put', 'dentro de este archivo', 'y luego escribiendo']):
                shared.session_state['stream_to_word'] = True
                if language == 'en':
                    response += "\n✍️ Writing mode enabled. Send next lines directly."
                else:
                    response += "\n✍️ Modo escritura activado. Envía las siguientes líneas directamente."
            return response
        
        elif any(word in text_lower for word in ['read', 'leu', 'leer', 'leí', 'lee', 'abr', 'open']):
            if not target or not location:
                command = shared.ai.understand_command(text, language)
                target = target or command.get('target')
                location = location or command.get('location')
            return await self._read_word_doc(target, location, language)
        
        elif any(word in text_lower for word in ['edit', 'editar', 'modif', 'cambiar', 'replace', 'overwrite', 'clear', 'borrar', 'reemplaz']):
            if not target or not location or not content:
                command = shared.ai.understand_command(text, language)
                target = target or command.get('target')
                location = location or command.get('location')
                content = content or command.get('content', '')
            if any(word in text_lower for word in ['replace', 'overwrite', 'clear', 'remove all', 'borrar', 'reemplaz', 'elimina todo']):
                return await self._replace_word_doc(target, location, content, language)
            return await self._edit_word_doc(target, location, content, language)
        
        # Default: If has target and location, likely a create operation
        if target and location:
            return await self._create_word_doc(target, location, content, language)
        
        return "Please specify create, read, or edit" if language == 'en' else "Especifica crear, leer o editar"

    async def append_to_last_word(self, raw_text: str, language: str) -> str:
        """Append a raw conversational line to the last active Word document."""
        try:
            full_path = self._shared_state.get('last_word_path')
            if not full_path:
                return "No active Word file found" if language == 'en' else "No se encontró un archivo Word activo"

            if not os.path.exists(full_path):
                return f"File not found: {full_path}" if language == 'en' else f"Archivo no encontrado: {full_path}"

            doc = Document(full_path)
            doc.add_paragraph(raw_text)
            doc.save(full_path)

            return "✓ Line added to current Word document" if language == 'en' else "✓ Línea agregada al documento Word actual"
        except PermissionError:
            return "Permission denied: close the file if it's open in Word" if language == 'en' else "Permiso denegado: cierra el archivo si está abierto en Word"
        except Exception as e:
            return f"Error appending to Word document: {str(e)}"
    
    async def _create_word_doc(self, filename: str, location: str, content: str, language: str) -> str:
        """Create a new Word document"""
        try:
            if not filename:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"

            if not location:
                location = os.path.join(os.path.expanduser('~'), 'Desktop')

            # Resolve location
            if location in ['desktop', 'escritorio']:
                location = os.path.join(os.path.expanduser('~'), 'Desktop')
            elif location in ['documents', 'documentos']:
                location = os.path.join(os.path.expanduser('~'), 'Documents')
            elif location in ['downloads', 'descargas']:
                location = os.path.join(os.path.expanduser('~'), 'Downloads')
            
            # Create full path
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            full_path = os.path.join(location, filename)
            
            # Create directory if needed
            os.makedirs(location, exist_ok=True)
            
            # Check if file exists first
            if os.path.exists(full_path):
                try:
                    test_doc = Document(full_path)
                except Exception as read_err:
                    msg = f"File exists but is corrupted: {full_path}" if language == 'en' else f"El archivo existe pero está corrupto: {full_path}"
                    return msg
            
            # Create document
            try:
                doc = Document()
                doc.add_heading('Document', level=0)
                
                if content:
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph('This document was created by JARVIS.')
                
                # Save with error handling
                doc.save(full_path)
            except PermissionError:
                return f"❌ Permission denied: {full_path}\n💡 Close the file if it's open in Word or another app." if language == 'en' else f"❌ Permiso denegado: {full_path}\n💡 Cierra el archivo si está abierto en Word u otra aplicación."
            
            # Calculate file size
            file_size = os.path.getsize(full_path)
            
            msg = f"✓ Word document created: {full_path}\n📄 Size: {file_size} bytes"
            if language == 'es':
                msg = f"✓ Documento Word creado: {full_path}\n📄 Tamaño: {file_size} bytes"
            
            try:
                self._shared_state['last_word_path'] = full_path
                self._shared_state['last_read_type'] = 'word'
            except Exception:
                pass
            
            return msg
            
        except Exception as e:
            error_msg = f"Error creating Word document: {str(e)}\n💡 Tip: Close the file if it's open in Word."
            if language == 'es':
                error_msg = f"Error al crear documento Word: {str(e)}\n💡 Consejo: Cierra el archivo si está abierto en Word."
            return error_msg
    
    async def _read_word_doc(self, filename: str, location: str, language: str) -> str:
        """Read content from Word document"""
        try:
            full_path = self._resolve_existing_word_path(filename, location)
            if not full_path:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"
            
            if not os.path.exists(full_path):
                return f"File not found: {full_path}" if language == 'en' else f"Archivo no encontrado: {full_path}"
            
            # Read document
            doc = Document(full_path)
            
            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            content = '\n'.join(full_text)
            
            # Preview (first 500 characters)
            preview = content[:500] + "..." if len(content) > 500 else content

            try:
                self._shared_state['last_word_path'] = full_path
            except Exception:
                pass
            
            msg = f"📄 Word Document: {filename}\n\n{preview}\n\n[Total: {len(content)} characters]"
            if language == 'es':
                msg = f"📄 Documento Word: {filename}\n\n{preview}\n\n[Total: {len(content)} caracteres]"
            
            return msg
            
        except Exception as e:
            return f"Error reading Word document: {str(e)}"
    
    async def _edit_word_doc(self, filename: str, location: str, new_content: str, language: str) -> str:
        """Edit existing Word document (append content)"""
        try:
            full_path = self._resolve_existing_word_path(filename, location)
            if not full_path:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"
            
            if not os.path.exists(full_path):
                return f"File not found: {full_path}" if language == 'en' else f"Archivo no encontrado: {full_path}"
            
            # Open and edit
            doc = Document(full_path)
            doc.add_paragraph(new_content)
            doc.save(full_path)

            try:
                self._shared_state['last_word_path'] = full_path
            except Exception:
                pass
            
            msg = f"✓ Word document edited: {full_path}"
            if language == 'es':
                msg = f"✓ Documento Word editado: {full_path}"
            
            return msg
            
        except Exception as e:
            return f"Error editing Word document: {str(e)}"

    async def _replace_word_doc(self, filename: str, location: str, new_content: str, language: str) -> str:
        """Replace all content in an existing Word document"""
        try:
            full_path = self._resolve_existing_word_path(filename, location)
            if not full_path:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"

            if not os.path.exists(full_path):
                return f"File not found: {full_path}" if language == 'en' else f"Archivo no encontrado: {full_path}"

            try:
                doc = Document(full_path)
            except PermissionError:
                return f"❌ Permission denied: {full_path}\n💡 Close the file if it's open in Word or another app." if language == 'en' else f"❌ Permiso denegado: {full_path}\n💡 Cierra el archivo si está abierto en Word u otra aplicación."
            except Exception:
                return f"File is corrupted or cannot be read: {full_path}" if language == 'en' else f"El archivo está corrupto o no se puede leer: {full_path}"

            for para in list(doc.paragraphs):
                p = para._element
                p.getparent().remove(p)
            doc.add_paragraph(new_content)
            
            try:
                doc.save(full_path)
            except PermissionError:
                return f"❌ Permission denied when saving: {full_path}\n💡 Close the file if it's open." if language == 'en' else f"❌ Permiso denegado al guardar: {full_path}\n💡 Cierra el archivo si está abierto."

            try:
                self._shared_state['last_word_path'] = full_path
            except Exception:
                pass

            msg = f"✓ Word document updated: {full_path}"
            if language == 'es':
                msg = f"✓ Documento Word actualizado: {full_path}"
            return msg

        except Exception as e:
            return f"Error editing Word document: {str(e)}"

    async def _summarize_word_doc(self, filename: str, location: str, output_name: str, language: str, shared) -> str:
        """Summarize a Word document and save to a new Word file"""
        try:
            full_path = self._resolve_existing_word_path(filename, location)
            if not full_path:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"

            if not os.path.exists(full_path):
                return f"File not found: {full_path}" if language == 'en' else f"Archivo no encontrado: {full_path}"

            doc = Document(full_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            if not paragraphs:
                return "The document is empty" if language == 'en' else "El documento esta vacio"

            chunks = self._chunk_paragraphs(paragraphs, max_chars=6000)
            summary = await self._summarize_chunks(chunks, language, shared)

            if output_name:
                output_filename = output_name
            else:
                base = os.path.splitext(os.path.basename(full_path))[0]
                output_filename = f"{base}_summary"

            if not output_filename.lower().endswith('.docx'):
                output_filename += '.docx'

            output_dir = os.path.dirname(full_path)
            output_path = os.path.join(output_dir, output_filename)

            out_doc = Document()
            out_doc.add_heading('Summary', level=1)
            out_doc.add_paragraph(summary)
            out_doc.save(output_path)

            try:
                self._shared_state['last_word_path'] = output_path
            except Exception:
                pass

            msg = f"✓ Summary created: {output_path}"
            if language == 'es':
                msg = f"✓ Resumen creado: {output_path}"
            return msg

        except Exception as e:
            return f"Error summarizing Word document: {str(e)}"

    def _chunk_paragraphs(self, paragraphs, max_chars=6000):
        """Group paragraphs into chunks sized for LLM summarization."""
        chunks = []
        current = []
        current_len = 0

        for para in paragraphs:
            if current_len + len(para) + 1 > max_chars and current:
                chunks.append("\n".join(current))
                current = [para]
                current_len = len(para)
            else:
                current.append(para)
                current_len += len(para) + 1

        if current:
            chunks.append("\n".join(current))

        return chunks

    async def _summarize_chunks(self, chunks, language, shared):
        """Summarize chunks and reduce into a final summary."""
        if not chunks:
            return ""

        section_prompt = "Summarize this section in 4-6 bullet points." if language == 'en' else "Resume esta seccion en 4-6 puntos."
        reduce_prompt = "Combine and refine these summaries into 6-10 bullet points." if language == 'en' else "Combina y refina estos resumenes en 6-10 puntos."

        partials = []
        for chunk in chunks:
            response = shared.ai.client.chat.completions.create(
                model=shared.ai.model,
                messages=[
                    {"role": "system", "content": section_prompt},
                    {"role": "user", "content": chunk}
                ],
                temperature=0.3,
                max_tokens=220,
            )
            partials.append(response.choices[0].message.content.strip())

        if len(partials) == 1:
            return partials[0]

        combined = "\n".join(partials)
        response = shared.ai.client.chat.completions.create(
            model=shared.ai.model,
            messages=[
                {"role": "system", "content": reduce_prompt},
                {"role": "user", "content": combined}
            ],
            temperature=0.3,
            max_tokens=260,
        )
        return response.choices[0].message.content.strip()

    def _resolve_existing_word_path(self, filename: str, location: str) -> str:
        """
        Resolve a word file path, searching common folders if location is missing.
        Falls back to fuzzy (partial-name) matching within the given folder so files
        with emoji prefixes or slightly different names are still found.
        """
        if not filename:
            return ""

        # Normalise extension
        stem = filename[:-5] if filename.lower().endswith('.docx') else filename
        filename_with_ext = stem + '.docx'

        # Resolve location shortcuts
        if location in ['desktop', 'escritorio']:
            location = os.path.join(os.path.expanduser('~'), 'Desktop')
        elif location in ['documents', 'documentos']:
            location = os.path.join(os.path.expanduser('~'), 'Documents')
        elif location in ['downloads', 'descargas']:
            location = os.path.join(os.path.expanduser('~'), 'Downloads')

        search_dirs = []
        if location:
            search_dirs.append(location)
        home = os.path.expanduser('~')
        for folder in ['Desktop', 'Documents', 'Downloads']:
            d = os.path.join(home, folder)
            if d not in search_dirs:
                search_dirs.append(d)

        for d in search_dirs:
            # 1) Exact match
            candidate = os.path.join(d, filename_with_ext)
            if os.path.exists(candidate):
                return candidate

            # 2) Fuzzy match: any .docx in that folder whose name contains the stem
            if os.path.isdir(d):
                stem_lower = stem.lower().strip()
                try:
                    for fname in os.listdir(d):
                        if fname.lower().endswith('.docx') and stem_lower in fname.lower():
                            return os.path.join(d, fname)
                except PermissionError:
                    pass

        # Nothing found – return the expected path so callers produce a clear error
        base = location if location else os.path.join(home, 'Desktop')
        return os.path.join(base, filename_with_ext)

    def _extract_word_details(self, text: str, language: str) -> Dict[str, str]:
        """Extract target/location/content from natural text without AI calls."""
        text_lower = text.lower()
        home = os.path.expanduser('~')

        location = None
        explicit_location = re.search(
            r'(?:in\s+(?:this\s+)?place|in\s+this\s+path|at\s+path|en\s+esta\s+ruta|en\s+este\s+lugar|ruta)\s+["\']([^"\']+)["\']',
            text,
            re.IGNORECASE,
        )
        if explicit_location:
            location = explicit_location.group(1).strip()

        if location is None and ('desktop' in text_lower or 'escritorio' in text_lower):
            location = os.path.join(home, 'Desktop')
        elif location is None and ('documents' in text_lower or 'documentos' in text_lower):
            location = os.path.join(home, 'Documents')
        elif location is None and ('downloads' in text_lower or 'descargas' in text_lower):
            location = os.path.join(home, 'Downloads')

        folder_match = re.search(
            r'(?:inside|into|in|within|dentro de|en)\s+(?:the\s+|la\s+)?(?:folder|carpeta)\s+["\']?([^"\'\n]+?)["\']?(?=\s+(?:called|named|name|con el nombre|and|with|write|escribe|guardar|save|$)|$)',
            text,
            re.IGNORECASE,
        )
        if folder_match and location:
            folder_name = folder_match.group(1).strip().strip('.,;:')
            if folder_name:
                location = os.path.join(location, folder_name)

        target = None
        target_match = re.search(
            r'(?:called|named|name|con el nombre|llamad[oa])\s+["\']?([^"\'\n]+?)["\']?(?=\s+(?:and|with|to|that|where|write|escribe|save|guardar|$)|$)',
            text,
            re.IGNORECASE,
        )
        if target_match:
            target = target_match.group(1).strip().strip('.,;:')

        content = ''
        quoted_blocks = re.findall(r'["“](.+?)["”]', text, re.DOTALL)
        if quoted_blocks:
            content = max((chunk.strip() for chunk in quoted_blocks if chunk.strip()), key=len, default='')
        else:
            content_match = re.search(
                r'(?:write|escribe|escribir|add|agrega|añade|with\s+(?:the\s+)?text)\s+(.*?)(?:\s+(?:and\s+save|save|y\s+guarda|guardar)\b|$)',
                text,
                re.IGNORECASE | re.DOTALL,
            )
            if content_match:
                content = content_match.group(1).strip().strip('"\'')

        if target and target.lower().endswith('.docx'):
            target = target[:-5]

        return {
            'target': target,
            'location': location,
            'content': content,
            'language': language,
        }

    def _extract_summary_details(self, text: str, language: str) -> Dict[str, str]:
        """Extract source and output names for summary commands."""
        details = {'source': None, 'output': None, 'location': None}

        # Explicit save/read path
        location_match = re.search(
            r'(?:in\s+(?:this\s+)?path|in\s+(?:this\s+)?place|save\s+in\s+(?:this\s+)?path|save\s+to\s+(?:this\s+)?path|ruta|en\s+esta\s+ruta)\s+["\']([^"\']+)["\']',
            text,
            re.IGNORECASE,
        )
        if location_match:
            details['location'] = location_match.group(1).strip()

        # Quoted names; first quoted item is usually source, second is often output name
        quoted_values = re.findall(r'["\']([^"\']+)["\']', text)
        filtered = [q.strip() for q in quoted_values if q and ':\\' not in q]

        source_match = re.search(
            r'(?:read|leer|summarize|summary|resumen|resumir)\s+(?:the\s+)?(?:file|archivo|documento|word)?\s*["\']([^"\']+?)(?:\.docx)?["\']',
            text,
            re.IGNORECASE,
        )
        if source_match:
            details['source'] = source_match.group(1).strip().strip('.,;:')
        elif filtered:
            details['source'] = filtered[0].strip().strip('.,;:')

        output_match = re.search(
            r'(?:save\s+as|guardar\s+como|named|called|con\s+el\s+nombre)\s+["\']([^"\']+?)(?:\.docx)?["\']',
            text,
            re.IGNORECASE,
        )
        if output_match:
            details['output'] = output_match.group(1).strip().strip('.,;:')
        elif len(filtered) >= 2:
            maybe_output = filtered[1].strip().strip('.,;:')
            if maybe_output.lower() not in ['a new document', 'new document', 'new file', 'in a new document']:
                details['output'] = maybe_output

        if details['source']:
            details['source'] = os.path.splitext(details['source'])[0]
        if details['output']:
            details['output'] = os.path.splitext(details['output'])[0]
        elif details['source']:
            details['output'] = f"{details['source']}_summary"

        return details
