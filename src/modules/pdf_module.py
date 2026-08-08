"""
PDF Reading Module - OpenClaw Pattern
Handles PDF file reading operations
"""
from typing import Dict, Any
import os
import re
from src.core.base_module import BaseModule

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class PDFModule(BaseModule):
    """Module for PDF reading operations"""
    
    def __init__(self, name: str):
        super().__init__(name)
        self.description = "Handles PDF file reading"
        
        if not PDF_AVAILABLE:
            print("⚠️ PyPDF2 not installed. Install with: pip install PyPDF2")
            self.enabled = False
    
    async def can_handle(self, intent: Dict[str, Any]) -> bool:
        """Check if this module handles PDFs"""
        return intent.get('type') == 'pdf' and self.enabled
    
    async def execute(self, intent: Dict[str, Any], shared) -> str:
        """Execute PDF reading operation"""
        text = intent.get('text', '')
        language = intent.get('language', 'en')

        self._shared_state = shared.session_state
        text_lower = text.lower()

        parsed = self._extract_pdf_details(text)
        parsed_target = parsed.get('target')
        parsed_location = parsed.get('location')
        parsed_output = parsed.get('output')
        
        if any(word in text_lower for word in ['summary', 'summarize', 'resumen', 'resumir', 'resume']):
            summary_details = self._extract_summary_details(text, language)
            source_name = summary_details.get('source') or parsed_target
            output_name = summary_details.get('output') or parsed_output
            summary_location = summary_details.get('location') or parsed_location

            if not source_name or not summary_location:
                command = shared.ai.understand_command(text, language)
                source_name = source_name or command.get('target')
                summary_location = summary_location or command.get('location')

            return await self._summarize_pdf(source_name, summary_location, output_name, language, shared)
        
        # Parse command using AI for regular read
        command = shared.ai.understand_command(text, language)
        target = parsed_target or command.get('target')
        location = parsed_location or command.get('location')

        if not target:
            last_path = shared.session_state.get('last_pdf_path')
            if last_path:
                target = os.path.splitext(os.path.basename(last_path))[0]
                location = os.path.dirname(last_path)

        return await self._read_pdf(target, location, language)
    
    async def _read_pdf(self, filename: str, location: str, language: str) -> str:
        """Read content from PDF file"""
        try:
            full_path = self._resolve_existing_pdf_path(filename, location)
            if not full_path:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"
            
            if not os.path.exists(full_path):
                return f"PDF not found: {full_path}" if language == 'en' else f"PDF no encontrado: {full_path}"
            
            # Read PDF
            with open(full_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extract text from all pages
                full_text = []
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    full_text.append(f"--- Page {page_num + 1} ---\n{text}")
                
                content = '\n\n'.join(full_text)
            
            # File size
            file_size = os.path.getsize(full_path)
            
            # Preview (first 800 characters)
            preview = content[:800] + "..." if len(content) > 800 else content

            try:
                self._shared_state['last_pdf_path'] = full_path
                self._shared_state['last_read_type'] = 'pdf'
            except Exception:
                pass
            
            msg = f"📕 PDF Document: {filename}\n"
            msg += f"📄 Pages: {num_pages} | Size: {file_size} bytes\n\n"
            msg += f"{preview}\n\n"
            msg += f"[Total: {len(content)} characters extracted]"
            
            if language == 'es':
                msg = f"📕 Documento PDF: {filename}\n"
                msg += f"📄 Páginas: {num_pages} | Tamaño: {file_size} bytes\n\n"
                msg += f"{preview}\n\n"
                msg += f"[Total: {len(content)} caracteres extraídos]"
            
            return msg
            
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    async def _summarize_pdf(self, filename: str, location: str, output_name: str, language: str, shared) -> str:
        """Summarize a PDF file and save to a new Word document"""
        try:
            if not DOCX_AVAILABLE:
                return "python-docx is not installed" if language == 'en' else "python-docx no esta instalado"

            full_path = self._resolve_existing_pdf_path(filename, location)
            if not full_path:
                return "Please provide a file name" if language == 'en' else "Por favor indica el nombre del archivo"

            if not os.path.exists(full_path):
                return f"PDF not found: {full_path}" if language == 'en' else f"PDF no encontrado: {full_path}"

            with open(full_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = []
                for page in pdf_reader.pages:
                    text = page.extract_text() or ""
                    if text.strip():
                        full_text.append(text.strip())

            if not full_text:
                return "The PDF is empty" if language == 'en' else "El PDF esta vacio"

            chunks = self._chunk_text(full_text, max_chars=6000)
            summary = await self._summarize_chunks(chunks, language, shared)

            # Use custom output name if provided, otherwise generate from source
            if output_name:
                if not output_name.lower().endswith('.docx'):
                    output_filename = f"{output_name}.docx"
                else:
                    output_filename = output_name
            else:
                base = os.path.splitext(os.path.basename(full_path))[0]
                output_filename = f"{base}_summary.docx"
            
            output_dir = location if location else os.path.dirname(full_path)
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, output_filename)

            out_doc = Document()
            out_doc.add_heading('Summary', level=1)
            out_doc.add_paragraph(summary)
            out_doc.save(output_path)

            try:
                self._shared_state['last_pdf_path'] = full_path
            except Exception:
                pass

            msg = f"✓ Summary created: {output_path}"
            if language == 'es':
                msg = f"✓ Resumen creado: {output_path}"
            return msg

        except Exception as e:
            return f"Error summarizing PDF: {str(e)}"

    def _extract_pdf_details(self, text: str) -> Dict[str, str]:
        """Deterministic extraction of PDF source filename, path, and output name."""
        details = {'target': None, 'location': None, 'output': None}

        path_match = re.search(
            r'(?:in\s+(?:this\s+)?path|in\s+(?:this\s+)?place|ruta|en\s+esta\s+ruta)\s+["\']([^"\']+)["\']',
            text,
            re.IGNORECASE,
        )
        if path_match:
            details['location'] = path_match.group(1).strip()

        quoted_values = re.findall(r'["\']([^"\']+)["\']', text)
        non_path = [q.strip() for q in quoted_values if q and ':\\' not in q]
        for item in non_path:
            if item.lower().endswith('.pdf'):
                details['target'] = os.path.splitext(item)[0]
                break

        if not details['target']:
            source_match = re.search(
                r'(?:pdf\s+called|called\s+that|archivo\s+pdf\s+llamado|named)\s+["\']?([^"\'\s,]+)',
                text,
                re.IGNORECASE,
            )
            if source_match:
                details['target'] = os.path.splitext(source_match.group(1).strip().strip('.,;:'))[0]

        output_match = re.search(
            r'(?:save\s+as|guardar\s+como|named|called|con\s+el\s+nombre)\s+["\']([^"\']+?)(?:\.docx)?["\']',
            text,
            re.IGNORECASE,
        )
        if output_match:
            details['output'] = os.path.splitext(output_match.group(1).strip().strip('.,;:'))[0]

        return details

    def _extract_summary_details(self, text: str, language: str) -> Dict[str, str]:
        """Extract source and output names for summary commands."""
        details = {'source': None, 'output': None, 'location': None}
        parsed = self._extract_pdf_details(text)
        details['source'] = parsed.get('target')
        details['output'] = parsed.get('output')
        details['location'] = parsed.get('location')

        if not details['output'] and details['source']:
            details['output'] = f"{details['source']}_summary"

        return details

    def _resolve_existing_pdf_path(self, filename: str, location: str) -> str:
        """
        Resolve a PDF path, searching common folders if location is missing.
        Falls back to fuzzy (partial-name) matching so files with emoji prefixes
        or minor name differences are still found.
        """
        if not filename:
            return ""

        stem = filename[:-4] if filename.lower().endswith('.pdf') else filename
        filename_with_ext = stem + '.pdf'

        if location in ['desktop', 'escritorio']:
            location = os.path.join(os.path.expanduser('~'), 'Desktop')
        elif location in ['documents', 'documentos']:
            location = os.path.join(os.path.expanduser('~'), 'Documents')
        elif location in ['downloads', 'descargas']:
            location = os.path.join(os.path.expanduser('~'), 'Downloads')

        search_dirs = []
        if location:
            search_dirs.append(location)
        home = os.path.expanduser('~')
        for folder in ['Desktop', 'Documents', 'Downloads']:
            d = os.path.join(home, folder)
            if d not in search_dirs:
                search_dirs.append(d)

        for d in search_dirs:
            # 1) Exact match
            candidate = os.path.join(d, filename_with_ext)
            if os.path.exists(candidate):
                return candidate

            # 2) Fuzzy: any .pdf in the folder containing the stem
            if os.path.isdir(d):
                stem_lower = stem.lower().strip()
                try:
                    for fname in os.listdir(d):
                        if fname.lower().endswith('.pdf') and stem_lower in fname.lower():
                            return os.path.join(d, fname)
                except PermissionError:
                    pass

        base = location if location else os.path.join(home, 'Desktop')
        return os.path.join(base, filename_with_ext)

    def _chunk_text(self, blocks, max_chars=6000):
        """Group text blocks into chunks sized for LLM summarization."""
        chunks = []
        current = []
        current_len = 0

        for block in blocks:
            if current_len + len(block) + 1 > max_chars and current:
                chunks.append("\n".join(current))
                current = [block]
                current_len = len(block)
            else:
                current.append(block)
                current_len += len(block) + 1

        if current:
            chunks.append("\n".join(current))

        return chunks

    async def _summarize_chunks(self, chunks, language, shared):
        """Summarize chunks and reduce into a final summary."""
        if not chunks:
            return ""

        section_prompt = "Summarize this section in 4-6 bullet points." if language == 'en' else "Resume esta seccion en 4-6 puntos."
        reduce_prompt = "Combine and refine these summaries into 6-10 bullet points." if language == 'en' else "Combina y refina estos resumenes en 6-10 puntos."

        partials = []
        for chunk in chunks:
            response = shared.ai.client.chat.completions.create(
                model=shared.ai.model,
                messages=[
                    {"role": "system", "content": section_prompt},
                    {"role": "user", "content": chunk}
                ],
                temperature=0.3,
                max_tokens=220,
            )
            partials.append(response.choices[0].message.content.strip())

        if len(partials) == 1:
            return partials[0]

        combined = "\n".join(partials)
        response = shared.ai.client.chat.completions.create(
            model=shared.ai.model,
            messages=[
                {"role": "system", "content": reduce_prompt},
                {"role": "user", "content": combined}
            ],
            temperature=0.3,
            max_tokens=260,
        )
        return response.choices[0].message.content.strip()

    def _extract_summary_details(self, text: str, language: str) -> Dict[str, str]:
        """Extract source/output/location for PDF summary commands."""
        details = {'source': None, 'output': None, 'location': None}

        location_match = re.search(
            r'(?:in\s+(?:this\s+)?path|in\s+(?:this\s+)?place|save\s+in\s+(?:this\s+)?path|ruta|en\s+esta\s+ruta)\s+["\']([^"\']+)["\']',
            text,
            re.IGNORECASE,
        )
        if location_match:
            details['location'] = location_match.group(1).strip()

        source_match = re.search(
            r'(?:pdf\s+called|archivo\s+pdf\s+llamado|called\s+that|named)\s+["\']?([^"\'\s,]+)(?:\.pdf)?["\']?',
            text,
            re.IGNORECASE,
        )
        if source_match:
            details['source'] = source_match.group(1).strip().strip('.,;:')
        else:
            quoted_values = re.findall(r'["\']([^"\']+)["\']', text)
            candidates = [q for q in quoted_values if ':\\' not in q and q.lower() not in ['a new document', 'new document']]
            if candidates:
                details['source'] = os.path.splitext(candidates[0])[0]

        output_match = re.search(
            r'(?:save\s+as|guardar\s+como|named|called|con\s+el\s+nombre)\s+["\']([^"\']+?)(?:\.docx)?["\']',
            text,
            re.IGNORECASE,
        )
        if output_match:
            details['output'] = os.path.splitext(output_match.group(1).strip().strip('.,;:'))[0]

        if not details['output'] and details['source']:
            details['output'] = f"{details['source']}_summary"

        return details
