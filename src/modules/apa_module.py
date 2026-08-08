"""
APA Module: Handles generation of APA-formatted thesis/report templates and documents.

Features:
- Create title pages with author, title, institution
- Generate table of contents
- Add standard APA sections (goals, recommendations, conclusion)
- Apply APA formatting (spacing, fonts, headings)
- Insert summarized content with proper citations
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Lazy imports to match OpenClaw pattern
Document = None
Pt = None
RGBColor = None
Inches = None

# Import base module
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))
from core.base_module import BaseModule


def _ensure_docx_imports():
    """Lazy-load docx dependencies."""
    global Document, Pt, RGBColor, Inches
    if Document is None:
        from docx import Document as DocxDocument
        from docx.shared import Pt as DxPt, RGBColor as DxRGBColor, Inches as DxInches
        Document = DocxDocument
        Pt = DxPt
        RGBColor = DxRGBColor
        Inches = DxInches


class APAModule(BaseModule):
    """Generates APA-formatted thesis templates and reports."""

    def __init__(self, name: str = "apa"):
        """Initialize APA module."""
        super().__init__(name)
        self.description = "Generate APA-formatted report templates and manage sections"

    async def can_handle(self, intent: Dict) -> bool:
        """Check if this module can handle APA-related requests."""
        intent_type = intent.get('type', '').lower()
        return intent_type == 'apa'

    async def execute(self, intent: Dict, shared) -> str:
        """
        Execute APA-related requests.
        
        Supports:
        - "create apa report template" → generates blank template
        - "create apa report for [title]" → generates template with title
        - "add goals to [doc]" → adds goals section to existing doc
        - "add recommendations to [doc]" → adds recommendations section
        """
        _ensure_docx_imports()
        
        user_input = intent.get('text', '')
        user_lower = user_input.lower()
        
        # Extract details from user input
        title, output_path, sections = self._extract_apa_details(user_input)
        
        # Store reference to shared resources for session state
        self.shared = shared
        
        # Determine action
        if "create" in user_lower and "template" in user_lower:
            return self._create_template(title, output_path, sections)
        elif "add" in user_lower and any(sec in user_lower for sec in ["goal", "recommendation", "conclusion", "abstract"]):
            return self._add_section_to_doc(user_input, title, output_path)
        elif "apa" in user_lower and "report" in user_lower:
            return self._create_template(title, output_path, ["goals", "recommendations", "conclusion"])
        else:
            return "I can help you create an APA report template. Try: 'create APA report template for [your title]' or 'add goals section to [document path]'"

    def _extract_apa_details(self, user_input: str) -> Tuple[str, str, List[str]]:
        """
        Extract title, output path, and sections from user input.
        
        Returns: (title, output_path, sections_list)
        """
        title = ""
        output_path = ""
        sections = ["goals", "recommendations", "conclusion"]
        
        # Extract quoted title
        quote_matches = re.findall(r'["\']([^"\']+)["\']', user_input)
        if quote_matches:
            title = quote_matches[0]
        
        # Extract explicit path
        path_match = re.search(r'(?:path|folder|to|in)\s+["\']?([^"\'\s]+(?:\\[^"\'\s]+)*)["\']?', user_input)
        if path_match:
            output_path = path_match.group(1)
        
        # Check for specific sections requested
        user_lower = user_input.lower()
        if "abstract" in user_lower:
            sections.insert(0, "abstract")
        if "literature review" in user_lower or "literature" in user_lower:
            sections.insert(1, "literature_review")
        if "methodology" in user_lower or "methods" in user_lower:
            sections.insert(2, "methodology")
        if "results" in user_lower:
            sections.insert(3, "results")
        if "discussion" in user_lower:
            sections.append("discussion")
        
        return title, output_path, sections

    def _create_template(self, title: str = "", output_path: str = "", sections: List[str] = None) -> str:
        """Create a new APA-formatted document template."""
        _ensure_docx_imports()
        
        if sections is None:
            sections = ["goals", "recommendations", "conclusion"]
        
        # Create document
        doc = Document()
        
        # Set default margins and font
        sections_obj = doc.sections
        for section in sections_obj:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title page
        self._add_apa_title_page(doc, title if title else "Untitled Report")
        
        # Add page break
        doc.add_page_break()
        
        # Add table of contents
        self._add_table_of_contents(doc, sections)
        
        # Add sections
        for section_name in sections:
            doc.add_page_break()
            self._add_apa_section(doc, section_name)
        
        # Determine output filename
        if not output_path:
            safe_title = re.sub(r'[\\/:*?"<>|]', '', title if title else "APA_Report")
            output_path = safe_title + "_APA.docx"
        elif not output_path.endswith('.docx'):
            output_path += ".docx"
        
        # Create directory if needed
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Save document
        doc.save(output_path)
        
        # Update session state if available
        if hasattr(self, 'shared') and self.shared:
            self.shared.session_state['last_word_path'] = output_path
            self.shared.save_session_state()
        
        return f"✅ APA report template created: {output_path}\n\nTemplate includes:\n- APA-formatted title page\n- Table of contents\n- Section stubs for: {', '.join(sections)}\n\nYou can now edit this file and add your content to each section."

    def _add_apa_title_page(self, doc: "Document", title: str) -> None:
        """Add APA-formatted title page to document."""
        _ensure_docx_imports()
        
        # Center alignment for title page
        paragraph = doc.add_paragraph()
        paragraph.alignment = 1  # Center alignment
        
        # Add blank lines for top spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Add title
        title_para = doc.add_paragraph(title.upper())
        title_para.alignment = 1  # Center
        title_format = title_para.paragraph_format
        title_format.line_spacing = 2.0  # Double space
        title_run = title_para.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Add blank lines
        for _ in range(3):
            doc.add_paragraph()
        
        # Add author line (placeholder)
        author_para = doc.add_paragraph("By Author Name")
        author_para.alignment = 1
        author_format = author_para.paragraph_format
        author_format.line_spacing = 2.0
        author_run = author_para.runs[0]
        author_run.font.size = Pt(12)
        
        # Add blank lines
        for _ in range(3):
            doc.add_paragraph()
        
        # Add institution/date
        institution_para = doc.add_paragraph("Institution Name\nUniversity")
        institution_para.alignment = 1
        institution_format = institution_para.paragraph_format
        institution_format.line_spacing = 2.0
        institution_run = institution_para.runs[0]
        institution_run.font.size = Pt(12)
        
        # Add date at bottom
        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.alignment = 1
        date_format = date_para.paragraph_format
        date_format.line_spacing = 2.0
        date_run = date_para.runs[0]
        date_run.font.size = Pt(12)

    def _add_table_of_contents(self, doc: "Document", sections: List[str]) -> None:
        """Add table of contents to document."""
        _ensure_docx_imports()
        
        # TOC heading
        toc_heading = doc.add_paragraph("TABLE OF CONTENTS")
        toc_heading_format = toc_heading.paragraph_format
        toc_heading_format.line_spacing = 2.0
        toc_heading_run = toc_heading.runs[0]
        toc_heading_run.font.bold = True
        toc_heading_run.font.size = Pt(12)
        
        doc.add_paragraph()  # Blank line
        
        # Add entries for each section
        for idx, section_name in enumerate(sections, 1):
            display_name = self._format_section_name(section_name)
            toc_entry = doc.add_paragraph(f"{display_name}........................................")
            toc_entry.paragraph_format.line_spacing = 2.0
            toc_entry_run = toc_entry.runs[0]
            toc_entry_run.font.size = Pt(12)

    def _add_apa_section(self, doc: "Document", section_name: str) -> None:
        """Add an APA-formatted section with heading and placeholder text."""
        _ensure_docx_imports()
        
        section_display = self._format_section_name(section_name)
        
        # Add section heading (centered, bold, level 1)
        heading = doc.add_paragraph(section_display)
        heading.alignment = 1  # Center
        heading_format = heading.paragraph_format
        heading_format.line_spacing = 2.0
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(12)
        heading_run = heading.runs[0]
        heading_run.font.bold = True
        heading_run.font.size = Pt(12)
        
        # Add placeholder text
        placeholder = doc.add_paragraph(f"[Insert {section_display.lower()} content here. Use APA style formatting with 12pt Times New Roman font and double spacing.]")
        placeholder_format = placeholder.paragraph_format
        placeholder_format.line_spacing = 2.0
        placeholder_format.first_line_indent = Inches(0.5)
        for run in placeholder.runs:
            run.font.size = Pt(12)
            run.font.italic = True

    def _add_section_to_doc(self, user_input: str, title: str = "", output_path: str = "") -> str:
        """Add a section to an existing document."""
        _ensure_docx_imports()
        
        # Extract section type from user input
        section_type = ""
        user_lower = user_input.lower()
        if "goal" in user_lower:
            section_type = "goals"
        elif "recommendation" in user_lower:
            section_type = "recommendations"
        elif "conclusion" in user_lower:
            section_type = "conclusion"
        elif "abstract" in user_lower:
            section_type = "abstract"
        elif "literature" in user_lower:
            section_type = "literature_review"
        elif "methodology" in user_lower or "method" in user_lower:
            section_type = "methodology"
        elif "result" in user_lower:
            section_type = "results"
        elif "discussion" in user_lower:
            section_type = "discussion"
        
        if not section_type:
            return "Could not identify section type. Try: 'add goals section to [document]' or 'add recommendations to [document]'"
        
        # Try to find document to edit
        doc_path = self._find_document_to_edit(user_input)
        if not doc_path or not os.path.exists(doc_path):
            # Try using session state
            if hasattr(self, 'shared') and self.shared and self.shared.session_state.get('last_word_path'):
                doc_path = self.shared.session_state['last_word_path']
            else:
                return f"Could not find document to edit. Please specify the path: 'add {section_type} to [document path]'"
        
        # Open document and add section
        try:
            doc = Document(doc_path)
        except Exception as e:
            return f"Error opening document {doc_path}: {str(e)}"
        
        # Add page break and section
        doc.add_page_break()
        self._add_apa_section(doc, section_type)
        
        # Save changes
        doc.save(doc_path)
        
        return f"✅ {self._format_section_name(section_type)} section added to {doc_path}"

    def _find_document_to_edit(self, user_input: str) -> Optional[str]:
        """Extract document path from user input."""
        # Look for quoted path
        quote_matches = re.findall(r'["\']([^"\']*\.docx)["\']', user_input, re.IGNORECASE)
        if quote_matches:
            return quote_matches[0]
        
        # Look for path pattern
        path_match = re.search(r'(?:to|edit|in|file|document)\s+["\']?([^"\'\s]+\.docx)["\']?', user_input, re.IGNORECASE)
        if path_match:
            return path_match.group(1)
        
        # Look for bare filename
        bare_match = re.search(r'([A-Za-z0-9_\-]+\.docx)', user_input)
        if bare_match:
            return bare_match.group(1)
        
        return None

    def _format_section_name(self, section_key: str) -> str:
        """Convert section key to display name."""
        mapping = {
            "abstract": "Abstract",
            "introduction": "Introduction",
            "literature_review": "Literature Review",
            "methodology": "Methodology",
            "results": "Results",
            "discussion": "Discussion",
            "goals": "Goals and Objectives",
            "recommendations": "Recommendations",
            "conclusion": "Conclusion",
            "references": "References",
        }
        return mapping.get(section_key.lower(), section_key.title())
