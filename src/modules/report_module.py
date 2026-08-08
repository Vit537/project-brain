"""
Report Module — SW Development Report Specialist
Creates, edits, and analyses academic engineering project reports (.docx)
following SCRUM + PUDS methodologies with proper Word formatting.

Capabilities:
- Cover page with university data, team members, dates
- Heading hierarchy (Heading 1 / 2 / 3) matching the standard report structure
- User story tables (HU cards)
- Sprint section templates (all 10 sub-sections)
- Daily scrum tables, backlog tables
- Report gap analysis (which sections are done / missing / incomplete)
- Table of Contents placeholder
"""

import json
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── Project config path ──────────────────────────────────────────────
CONFIG_PATH = os.path.abspath(
    os.path.join(os.path.dirname(__file__), '..', '..', 'config', 'report_project.json')
)

# ── Standard report structure ────────────────────────────────────────
# Each tuple: (heading_level, section_number, title_es, title_en)
REPORT_SKELETON = [
    (1, "1",   "Introducción",                        "Introduction"),
    (1, "2",   "Antecedentes (Marco Teórico)",        "Background (Theoretical Framework)"),
    (2, "2.1", "Historia",                             "History"),
    (2, "2.2", "Software Existente",                   "Existing Software"),
    (2, "2.3", "Casos de Estudio",                     "Case Studies"),
    (1, "3",   "Descripción del Problema",             "Problem Description"),
    (2, "3.1", "Planteamiento del Problema",           "Problem Statement"),
    (2, "3.2", "Formulación del Problema",             "Problem Formulation"),
    (1, "4",   "Objetivos",                            "Objectives"),
    (2, "4.1", "Objetivo General",                     "General Objective"),
    (2, "4.2", "Objetivos Específicos",                "Specific Objectives"),
    (1, "5",   "Alcance",                              "Scope"),
    (1, "6",   "Elementos del CBIS",                   "CBIS Elements"),
    (1, "7",   "Tecnología",                           "Technology"),
    (1, "8",   "Costos",                               "Costs"),
    (1, "9",   "Beneficios",                           "Benefits"),
    (1, "10",  "Sprints",                              "Sprints"),
    (1, "11",  "Bibliografía",                         "Bibliography"),
    (1, "12",  "Anexos",                               "Annexes"),
]

# Sub-sections generated for every Sprint N (N >= 0)
SPRINT_0_SECTIONS = [
    (2, ".1", "Equipo Scrum",                          "Scrum Team"),
    (2, ".2", "Objetivo del Producto",                 "Product Goal"),
    (2, ".3", "Requerimientos Iniciales",              "Initial Requirements"),
    (2, ".4", "Duración del Sprint",                   "Sprint Duration"),
    (2, ".5", "Infraestructura Tecnológica",           "Technology Infrastructure"),
    (2, ".6", "Patrón de Desarrollo",                  "Development Pattern"),
    (2, ".7", "Modelos Iniciales (UML)",               "Initial Models (UML)"),
    (2, ".8", "Criterios de Calidad",                  "Quality Criteria"),
    (2, ".9", "Product Backlog",                       "Product Backlog"),
]

SPRINT_N_SECTIONS = [
    (2, ".1", "Sprint Planning",                       "Sprint Planning"),
    (2, ".2", "Objetivos del Sprint",                  "Sprint Goals"),
    (2, ".3", "Equipo Scrum",                          "Scrum Team"),
    (2, ".4", "Story Points (Esfuerzo)",               "Story Points (Effort)"),
    (3, ".4.1", "Diagrama de Casos de Uso",            "Use Case Diagram"),
    (2, ".5", "Historias de Usuario (HU)",             "User Stories (HU)"),
    (2, ".6", "Diseños (Diagramas UML)",               "Designs (UML Diagrams)"),
    (3, ".6.1", "Tareas para Iteración del Sprint",    "Tasks for Sprint Iteration"),
    (2, ".7", "Sprint Backlog",                        "Sprint Backlog"),
    (2, ".8", "Daily Scrum",                           "Daily Scrum"),
    (2, ".9", "Sprint Review",                         "Sprint Review"),
    (2, ".10", "Sprint Retrospective",                 "Sprint Retrospective"),
]


class ReportModule:
    """
    Specialist module for SW development project reports.
    Works with python-docx to produce properly formatted .docx files.
    """

    def __init__(self):
        self._config: Optional[Dict[str, Any]] = None
        if not DOCX_AVAILABLE:
            print("⚠️  python-docx not installed — report module disabled")

    # ──────────────────────────────────────────────────────────────────
    # Project config helpers
    # ──────────────────────────────────────────────────────────────────
    def _load_config(self) -> Dict[str, Any]:
        """Load project config from JSON."""
        if self._config is not None:
            return self._config
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
                self._config = json.load(f)
        else:
            self._config = {}
        return self._config

    def _save_config(self, cfg: Dict[str, Any]) -> None:
        """Persist project config to disk."""
        os.makedirs(os.path.dirname(CONFIG_PATH), exist_ok=True)
        with open(CONFIG_PATH, 'w', encoding='utf-8') as f:
            json.dump(cfg, f, ensure_ascii=False, indent=4)
        self._config = cfg

    def _is_configured(self) -> bool:
        cfg = self._load_config()
        return cfg.get('configured', False)

    # ──────────────────────────────────────────────────────────────────
    # 1.  SETUP PROJECT
    # ──────────────────────────────────────────────────────────────────
    async def setup_project(self, data: Dict[str, Any], language: str = "es") -> str:
        """
        Store project metadata.  Accepts partial updates — merges into existing config.
        """
        cfg = self._load_config()

        # Merge scalar fields
        for key in [
            "university", "faculty", "subject", "professor",
            "group_number", "project_name", "semester", "year",
            "city", "country", "uml_version", "language",
            "sprint_duration_weeks",
        ]:
            if key in data and data[key]:
                cfg[key] = data[key]

        # Merge lists / dicts
        if "team_members" in data and data["team_members"]:
            cfg["team_members"] = data["team_members"]
        if "methodologies" in data and data["methodologies"]:
            cfg["methodologies"] = data["methodologies"]
        if "tech_stack" in data and data["tech_stack"]:
            if isinstance(data["tech_stack"], dict):
                cfg.setdefault("tech_stack", {}).update(data["tech_stack"])
            else:
                cfg["tech_stack"] = data["tech_stack"]
        if "story_point_scale" in data:
            cfg["story_point_scale"] = data["story_point_scale"]

        cfg["configured"] = True
        self._save_config(cfg)

        if language == "es":
            return (
                f"✅ Proyecto configurado: **{cfg.get('project_name', '(sin nombre)')}**\n"
                f"   Universidad: {cfg.get('university', '-')}\n"
                f"   Materia: {cfg.get('subject', '-')}\n"
                f"   Grupo: {cfg.get('group_number', '-')}\n"
                f"   Miembros: {len(cfg.get('team_members', []))}\n"
                f"   Metodologías: {', '.join(cfg.get('methodologies', []))}\n"
                f"   Archivo de config: {CONFIG_PATH}"
            )
        return (
            f"✅ Project configured: **{cfg.get('project_name', '(unnamed)')}**\n"
            f"   University: {cfg.get('university', '-')}\n"
            f"   Subject: {cfg.get('subject', '-')}\n"
            f"   Group: {cfg.get('group_number', '-')}\n"
            f"   Members: {len(cfg.get('team_members', []))}\n"
            f"   Methodologies: {', '.join(cfg.get('methodologies', []))}\n"
            f"   Config file: {CONFIG_PATH}"
        )

    # ──────────────────────────────────────────────────────────────────
    # 2.  CREATE REPORT TEMPLATE  (full skeleton)
    # ──────────────────────────────────────────────────────────────────
    async def create_report_template(
        self,
        filename: str,
        location: str,
        include_sprint_0: bool = True,
        language: str = "es",
    ) -> str:
        """Create a new .docx with cover page + all section headings."""
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        cfg = self._load_config()
        lang = language or cfg.get("language", "es")
        idx = 0 if lang == "es" else 1  # index into title tuples

        filename = self._ensure_docx(filename)
        full_path = self._resolve_path(filename, location)
        os.makedirs(os.path.dirname(full_path), exist_ok=True)

        doc = Document()
        self._set_default_style(doc)

        # ── Cover page ──
        self._build_cover_page(doc, cfg, lang)
        doc.add_page_break()

        # ── TOC placeholder ──
        self._add_toc_placeholder(doc, lang)
        doc.add_page_break()

        # ── Skeleton sections ──
        for level, num, title_es, title_en in REPORT_SKELETON:
            title = title_es if lang == "es" else title_en
            heading = f"{num}. {title}" if level == 1 else f"{num} {title}"
            doc.add_heading(heading, level=level)
            doc.add_paragraph("")  # empty body placeholder

        # ── Sprint 0 ──
        if include_sprint_0:
            self._add_sprint_headings(doc, 0, cfg, lang)

        doc.save(full_path)
        size = os.path.getsize(full_path)

        sections_count = len(REPORT_SKELETON) + (len(SPRINT_0_SECTIONS) if include_sprint_0 else 0)
        if lang == "es":
            return (
                f"✅ Plantilla de reporte creada: {full_path}\n"
                f"📄 {sections_count} secciones | {size} bytes\n"
                f"Incluye: Portada, Índice (placeholder), todas las secciones estándar"
                + (", Sprint 0" if include_sprint_0 else "")
            )
        return (
            f"✅ Report template created: {full_path}\n"
            f"📄 {sections_count} sections | {size} bytes\n"
            f"Includes: Cover page, TOC placeholder, all standard sections"
            + (", Sprint 0" if include_sprint_0 else "")
        )

    # ──────────────────────────────────────────────────────────────────
    # 3.  WRITE REPORT SECTION
    # ──────────────────────────────────────────────────────────────────
    async def write_section(
        self,
        filename: str,
        location: str,
        section_title: str,
        content: str,
        language: str = "es",
    ) -> str:
        """
        Write / overwrite the body content under a specific heading.
        Finds the heading in the document and replaces everything between
        it and the next heading of equal or higher level.
        """
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)
        heading_idx, heading_level = self._find_heading(doc, section_title)

        if heading_idx is None:
            return (
                f"❌ Sección '{section_title}' no encontrada en el documento."
                if language == "es" else
                f"❌ Section '{section_title}' not found in document."
            )

        # Remove existing body paragraphs under that heading
        self._clear_section_body(doc, heading_idx, heading_level)

        # Insert new paragraphs after the heading
        insert_after = doc.paragraphs[heading_idx]
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Support sub-heading markers: lines starting with ## or ###
            if line.startswith('### '):
                new_p = self._insert_paragraph_after(insert_after, line[4:], heading_level=3)
            elif line.startswith('## '):
                new_p = self._insert_paragraph_after(insert_after, line[3:], heading_level=2)
            else:
                new_p = self._insert_paragraph_after(insert_after, line)
            insert_after = new_p

        doc.save(full_path)
        if language == "es":
            return f"✅ Sección '{section_title}' actualizada en {full_path}"
        return f"✅ Section '{section_title}' updated in {full_path}"

    # ──────────────────────────────────────────────────────────────────
    # 4.  CREATE / UPDATE COVER PAGE
    # ──────────────────────────────────────────────────────────────────
    async def create_cover_page(
        self,
        filename: str,
        location: str,
        language: str = "es",
    ) -> str:
        """Create or overwrite the cover page of an existing report."""
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        cfg = self._load_config()
        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)

        # Remove everything before the first Heading 1 (that's the old cover page)
        first_h1 = None
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading') and self._heading_level(p) == 1:
                first_h1 = i
                break

        if first_h1 is not None and first_h1 > 0:
            for _ in range(first_h1):
                elem = doc.paragraphs[0]._element
                elem.getparent().remove(elem)

        # Insert cover page elements BEFORE the first remaining paragraph
        body = doc.element.body
        first_child = body[0] if len(body) > 0 else None
        cover_elements = self._make_cover_elements(doc, cfg, language)
        for el in reversed(cover_elements):
            if first_child is not None:
                first_child.addprevious(el)
            else:
                body.append(el)

        doc.save(full_path)
        if language == "es":
            return f"✅ Portada actualizada en {full_path}"
        return f"✅ Cover page updated in {full_path}"

    # ──────────────────────────────────────────────────────────────────
    # 5.  ADD SPRINT SECTION
    # ──────────────────────────────────────────────────────────────────
    async def add_sprint(
        self,
        filename: str,
        location: str,
        sprint_number: int,
        sprint_goal: str = "",
        language: str = "es",
    ) -> str:
        """Add a complete Sprint N section with all sub-headings."""
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        cfg = self._load_config()
        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)

        # Find the "Sprints" heading (section 10) or the Bibliography heading to insert before
        insert_before_idx = None
        for i, p in enumerate(doc.paragraphs):
            text_lower = p.text.lower().strip()
            if '11' in text_lower and ('bibliograf' in text_lower or 'bibliography' in text_lower):
                insert_before_idx = i
                break

        self._add_sprint_headings(doc, sprint_number, cfg, language, insert_before_idx)

        # If a sprint goal was provided, write it under "Sprint Goals"
        if sprint_goal:
            goal_section = (
                f"10.{sprint_number}.2 Objetivos del Sprint"
                if language == "es" else
                f"10.{sprint_number}.2 Sprint Goals"
            )
            idx, _ = self._find_heading(doc, goal_section)
            if idx is not None:
                insert_after = doc.paragraphs[idx]
                self._insert_paragraph_after(insert_after, sprint_goal)

        doc.save(full_path)

        sub_count = len(SPRINT_0_SECTIONS) if sprint_number == 0 else len(SPRINT_N_SECTIONS)
        if language == "es":
            return f"✅ Sprint {sprint_number} agregado a {full_path} ({sub_count} sub-secciones)"
        return f"✅ Sprint {sprint_number} added to {full_path} ({sub_count} sub-sections)"

    # ──────────────────────────────────────────────────────────────────
    # 6.  CREATE USER STORY TABLE
    # ──────────────────────────────────────────────────────────────────
    async def create_user_story_table(
        self,
        filename: str,
        location: str,
        stories: List[Dict[str, Any]],
        section_title: str = "",
        language: str = "es",
    ) -> str:
        """
        Insert formatted User Story cards as tables.

        Each story dict should have:
          id, name, priority, story_points, status,
          as_role, i_want, so_that, acceptance_criteria (list), developer
        """
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)

        # Find insertion point
        insert_after_p = None
        if section_title:
            idx, _ = self._find_heading(doc, section_title)
            if idx is not None:
                insert_after_p = doc.paragraphs[idx]

        if insert_after_p is None:
            insert_after_p = doc.paragraphs[-1]

        for story in stories:
            insert_after_p = self._insert_hu_card(doc, insert_after_p, story, language)

        doc.save(full_path)

        if language == "es":
            return f"✅ {len(stories)} historia(s) de usuario insertada(s) en {full_path}"
        return f"✅ {len(stories)} user story/stories inserted in {full_path}"

    # ──────────────────────────────────────────────────────────────────
    # 7.  ADD DAILY SCRUM TABLE
    # ──────────────────────────────────────────────────────────────────
    async def add_daily_scrum_table(
        self,
        filename: str,
        location: str,
        sprint_number: int,
        entries: List[Dict[str, str]],
        language: str = "es",
    ) -> str:
        """
        Insert a Daily Scrum table.

        Each entry dict: {date, member, did_yesterday, doing_today, blockers}
        """
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)

        # Locate the Daily Scrum heading for this sprint
        target = f"10.{sprint_number}.8"
        idx, _ = self._find_heading(doc, target)
        insert_after_p = doc.paragraphs[idx] if idx is not None else doc.paragraphs[-1]

        # Build table
        headers_es = ["Fecha", "Miembro", "¿Qué hizo ayer?", "¿Qué hará hoy?", "Impedimentos"]
        headers_en = ["Date", "Member", "Did Yesterday", "Doing Today", "Blockers"]
        headers = headers_es if language == "es" else headers_en

        table = self._insert_table_after(doc, insert_after_p, headers, [
            [
                e.get("date", ""),
                e.get("member", ""),
                e.get("did_yesterday", ""),
                e.get("doing_today", ""),
                e.get("blockers", "—"),
            ]
            for e in entries
        ])

        doc.save(full_path)
        if language == "es":
            return f"✅ Tabla de Daily Scrum (Sprint {sprint_number}) con {len(entries)} registros insertada en {full_path}"
        return f"✅ Daily Scrum table (Sprint {sprint_number}) with {len(entries)} entries inserted in {full_path}"

    # ──────────────────────────────────────────────────────────────────
    # 8.  ADD SPRINT BACKLOG TABLE
    # ──────────────────────────────────────────────────────────────────
    async def add_backlog_table(
        self,
        filename: str,
        location: str,
        sprint_number: int,
        items: List[Dict[str, str]],
        is_product_backlog: bool = False,
        language: str = "es",
    ) -> str:
        """
        Insert a Sprint Backlog or Product Backlog table.

        Each item dict: {id, description, priority, story_points, status, responsible}
        """
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)

        # Find heading
        if is_product_backlog:
            target = f"10.{sprint_number}.9" if sprint_number == 0 else f"10.{sprint_number}.7"
        else:
            target = f"10.{sprint_number}.7"
        idx, _ = self._find_heading(doc, target)
        insert_after_p = doc.paragraphs[idx] if idx is not None else doc.paragraphs[-1]

        headers_es = ["ID", "Descripción", "Prioridad", "PHU", "Estado", "Responsable"]
        headers_en = ["ID", "Description", "Priority", "SP", "Status", "Responsible"]
        headers = headers_es if language == "es" else headers_en

        self._insert_table_after(doc, insert_after_p, headers, [
            [
                item.get("id", ""),
                item.get("description", ""),
                item.get("priority", ""),
                str(item.get("story_points", "")),
                item.get("status", ""),
                item.get("responsible", ""),
            ]
            for item in items
        ])

        doc.save(full_path)
        label = "Product Backlog" if is_product_backlog else "Sprint Backlog"
        if language == "es":
            return f"✅ Tabla de {label} (Sprint {sprint_number}) con {len(items)} elementos insertada en {full_path}"
        return f"✅ {label} table (Sprint {sprint_number}) with {len(items)} items inserted in {full_path}"

    # ──────────────────────────────────────────────────────────────────
    # 9.  ANALYSE REPORT PROGRESS
    # ──────────────────────────────────────────────────────────────────
    async def analyze_report(
        self,
        filename: str,
        location: str,
        language: str = "es",
    ) -> str:
        """
        Read the document, catalogue all headings, and report which
        sections are present (with content), present (empty), or missing.
        """
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)

        # Catalogue headings found in the document
        found_headings: Dict[str, str] = {}   # normalised title → "complete" | "empty"
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading'):
                title = p.text.strip()
                # Check if there's body text before the next heading
                has_body = False
                for j in range(i + 1, len(doc.paragraphs)):
                    nxt = doc.paragraphs[j]
                    if nxt.style.name.startswith('Heading'):
                        break
                    if nxt.text.strip():
                        has_body = True
                        break
                found_headings[self._normalise_heading(title)] = "complete" if has_body else "empty"

        # Also count tables
        table_count = len(doc.tables)

        # Compare against expected skeleton
        cfg = self._load_config()
        lang = language or cfg.get("language", "es")
        idx_t = 0 if lang == "es" else 1

        complete = []
        empty = []
        missing = []

        for level, num, title_es, title_en in REPORT_SKELETON:
            title = title_es if lang == "es" else title_en
            key = self._normalise_heading(f"{num} {title}")
            alt_key = self._normalise_heading(f"{num}. {title}")
            if key in found_headings or alt_key in found_headings:
                status = found_headings.get(key) or found_headings.get(alt_key)
                if status == "complete":
                    complete.append(f"  ✅ {num}. {title}")
                else:
                    empty.append(f"  ⬜ {num}. {title}")
            else:
                missing.append(f"  ❌ {num}. {title}")

        # Build report
        lines = []
        if lang == "es":
            lines.append(f"📊 **Análisis del reporte**: {full_path}")
            lines.append(f"   Párrafos: {len(doc.paragraphs)} | Tablas: {table_count}")
            lines.append(f"   Encabezados encontrados: {len(found_headings)}")
            lines.append("")
            if complete:
                lines.append("**Secciones con contenido:**")
                lines.extend(complete)
            if empty:
                lines.append("\n**Secciones sin contenido (solo título):**")
                lines.extend(empty)
            if missing:
                lines.append("\n**Secciones faltantes:**")
                lines.extend(missing)
            if not missing and not empty:
                lines.append("\n🎉 ¡Todas las secciones estándar tienen contenido!")
        else:
            lines.append(f"📊 **Report analysis**: {full_path}")
            lines.append(f"   Paragraphs: {len(doc.paragraphs)} | Tables: {table_count}")
            lines.append(f"   Headings found: {len(found_headings)}")
            lines.append("")
            if complete:
                lines.append("**Sections with content:**")
                lines.extend(complete)
            if empty:
                lines.append("\n**Sections without content (heading only):**")
                lines.extend(empty)
            if missing:
                lines.append("\n**Missing sections:**")
                lines.extend(missing)
            if not missing and not empty:
                lines.append("\n🎉 All standard sections have content!")

        return "\n".join(lines)

    # ──────────────────────────────────────────────────────────────────
    # 10. ADD SPRINT RETROSPECTIVE TABLE
    # ──────────────────────────────────────────────────────────────────
    async def add_retrospective(
        self,
        filename: str,
        location: str,
        sprint_number: int,
        went_well: List[str],
        went_wrong: List[str],
        to_improve: List[str],
        language: str = "es",
    ) -> str:
        """Insert a Sprint Retrospective table (What went well / wrong / improve)."""
        if not DOCX_AVAILABLE:
            return "❌ python-docx not installed"

        full_path = self._resolve_existing(filename, location)
        if not os.path.exists(full_path):
            return f"❌ Archivo no encontrado: {full_path}" if language == "es" else f"❌ File not found: {full_path}"

        doc = Document(full_path)
        target = f"10.{sprint_number}.10"
        idx, _ = self._find_heading(doc, target)
        insert_after_p = doc.paragraphs[idx] if idx is not None else doc.paragraphs[-1]

        if language == "es":
            headers = ["¿Qué salió bien?", "¿Qué salió mal?", "¿Qué mejorar?"]
        else:
            headers = ["What went well?", "What went wrong?", "What to improve?"]

        max_rows = max(len(went_well), len(went_wrong), len(to_improve))
        rows = []
        for i in range(max_rows):
            rows.append([
                went_well[i] if i < len(went_well) else "",
                went_wrong[i] if i < len(went_wrong) else "",
                to_improve[i] if i < len(to_improve) else "",
            ])

        self._insert_table_after(doc, insert_after_p, headers, rows)
        doc.save(full_path)

        if language == "es":
            return f"✅ Retrospectiva del Sprint {sprint_number} insertada en {full_path}"
        return f"✅ Sprint {sprint_number} Retrospective inserted in {full_path}"

    # ==================================================================
    #  PRIVATE HELPERS — Word formatting
    # ==================================================================

    def _ensure_docx(self, name: str) -> str:
        if not name.lower().endswith('.docx'):
            name += '.docx'
        return name

    def _resolve_path(self, filename: str, location: str) -> str:
        """Resolve location shortcuts to full path."""
        if not location:
            location = os.path.join(os.path.expanduser('~'), 'Desktop')
        shortcuts = {
            'desktop': os.path.join(os.path.expanduser('~'), 'Desktop'),
            'escritorio': os.path.join(os.path.expanduser('~'), 'Desktop'),
            'documents': os.path.join(os.path.expanduser('~'), 'Documents'),
            'documentos': os.path.join(os.path.expanduser('~'), 'Documents'),
            'downloads': os.path.join(os.path.expanduser('~'), 'Downloads'),
            'descargas': os.path.join(os.path.expanduser('~'), 'Downloads'),
        }
        location = shortcuts.get(location.lower().strip(), location)
        return os.path.join(location, filename)

    def _resolve_existing(self, filename: str, location: str) -> str:
        """Resolve an existing file with fuzzy matching."""
        filename = self._ensure_docx(filename)
        stem = os.path.splitext(filename)[0].lower().strip()

        if not location:
            location = os.path.join(os.path.expanduser('~'), 'Desktop')
        shortcuts = {
            'desktop': os.path.join(os.path.expanduser('~'), 'Desktop'),
            'escritorio': os.path.join(os.path.expanduser('~'), 'Desktop'),
            'documents': os.path.join(os.path.expanduser('~'), 'Documents'),
            'documentos': os.path.join(os.path.expanduser('~'), 'Documents'),
            'downloads': os.path.join(os.path.expanduser('~'), 'Downloads'),
            'descargas': os.path.join(os.path.expanduser('~'), 'Downloads'),
        }
        location = shortcuts.get(location.lower().strip(), location)

        # Exact
        candidate = os.path.join(location, filename)
        if os.path.exists(candidate):
            return candidate

        # Fuzzy within location
        if os.path.isdir(location):
            try:
                for fname in os.listdir(location):
                    if fname.lower().endswith('.docx') and stem in fname.lower():
                        return os.path.join(location, fname)
            except PermissionError:
                pass

        return candidate  # return expected path so caller gives clear error

    def _set_default_style(self, doc: "Document") -> None:
        """Set default font to Times New Roman 12pt, normal margins."""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.54)

    # ── Cover page building ──────────────────────────────────────────

    def _build_cover_page(self, doc: "Document", cfg: Dict, lang: str) -> None:
        """Add cover page paragraphs to a brand-new document."""
        university = cfg.get("university", "Universidad")
        faculty = cfg.get("faculty", "Facultad de Ingeniería")
        subject = cfg.get("subject", "")
        professor = cfg.get("professor", "")
        project_name = cfg.get("project_name", "Proyecto")
        group = cfg.get("group_number", "")
        team = cfg.get("team_members", [])
        year = cfg.get("year", str(datetime.now().year))
        semester = cfg.get("semester", "")
        city = cfg.get("city", "")
        country = cfg.get("country", "")

        def centered(text, bold=False, size=12, spacing_after=0):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            if spacing_after:
                p.paragraph_format.space_after = Pt(spacing_after)

        # Top block
        centered(university.upper(), bold=True, size=14, spacing_after=4)
        centered(faculty, bold=False, size=12, spacing_after=4)

        # Spacing
        doc.add_paragraph("")

        # Logo placeholder
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run("[LOGO DE LA UNIVERSIDAD]" if lang == "es" else "[UNIVERSITY LOGO]")
        run_logo.font.size = Pt(10)
        run_logo.font.color.rgb = RGBColor(150, 150, 150)

        doc.add_paragraph("")

        # Project label
        label = "PERFIL DE PROYECTO" if lang == "es" else "PROJECT PROFILE"
        centered(label, bold=True, size=16, spacing_after=8)

        # Project name
        centered(project_name.upper(), bold=True, size=14, spacing_after=12)

        # Subject
        if subject:
            lbl = "Materia" if lang == "es" else "Subject"
            centered(f"{lbl}: {subject}", size=12, spacing_after=4)
        if professor:
            lbl = "Docente" if lang == "es" else "Professor"
            centered(f"{lbl}: {professor}", size=12, spacing_after=4)
        if group:
            lbl = "Grupo" if lang == "es" else "Group"
            centered(f"{lbl}: {group}", size=12, spacing_after=8)

        # Team members
        if team:
            lbl = "Integrantes" if lang == "es" else "Team Members"
            centered(lbl, bold=True, size=12, spacing_after=4)
            for m in team:
                name = m.get("name", "")
                sid = m.get("student_id", "")
                line = f"{name}  —  {sid}" if sid else name
                centered(line, size=11, spacing_after=2)

        doc.add_paragraph("")

        # Date / location
        date_str = datetime.now().strftime("%d de %B de %Y") if lang == "es" else datetime.now().strftime("%B %d, %Y")
        location_str = f"{city}, {country}" if city and country else city or country or ""
        if location_str:
            centered(f"{location_str} — {date_str}", size=11)
        else:
            centered(date_str, size=11)

        # Academic year
        if semester or year:
            centered(f"{semester} — {year}" if semester else year, size=11)

    def _make_cover_elements(self, doc, cfg, lang):
        """Return a list of OxmlElement paragraphs for cover rewrite.
        (Used by create_cover_page to update an existing doc.)
        """
        # Create a temporary document, build cover, extract XML elements
        from docx import Document as TmpDoc
        tmp = TmpDoc()
        self._set_default_style(tmp)
        self._build_cover_page(tmp, cfg, lang)
        # also add a page break at the end
        tmp.add_page_break()
        elements = []
        for p in tmp.paragraphs:
            elements.append(p._element)
        return elements

    # ── TOC placeholder ──────────────────────────────────────────────

    def _add_toc_placeholder(self, doc, lang: str) -> None:
        """Insert a TOC field code that Word will refresh on open."""
        title = "CONTENIDO" if lang == "es" else "TABLE OF CONTENTS"
        doc.add_heading(title, level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = paragraph.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run2._r.append(instrText)
        run3 = paragraph.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)
        run4 = paragraph.add_run()
        hint = "Actualizar tabla de contenido (Ctrl+A, F9)" if lang == "es" else "Update table of contents (Ctrl+A, F9)"
        run4.add_text(hint)
        run4.font.color.rgb = RGBColor(150, 150, 150)
        run5 = paragraph.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)

    # ── Sprint heading generation ────────────────────────────────────

    def _add_sprint_headings(
        self,
        doc,
        sprint_number: int,
        cfg: Dict,
        lang: str,
        insert_before_idx: Optional[int] = None,
    ) -> None:
        """Add Sprint N heading + all sub-headings to the document.
        
        If insert_before_idx is given, the headings are inserted BEFORE that
        paragraph index (e.g. before Bibliography).  Otherwise they are appended
        at the end.
        """
        sections = SPRINT_0_SECTIONS if sprint_number == 0 else SPRINT_N_SECTIONS
        base_num = f"10.{sprint_number}"

        if insert_before_idx is not None and insert_before_idx < len(doc.paragraphs):
            # Insert before the target paragraph using XML manipulation
            ref_element = doc.paragraphs[insert_before_idx]._element
            elements_to_insert = []

            # Sprint heading
            heading_text = f"{base_num} Sprint {sprint_number}"
            elements_to_insert.append(('heading', heading_text, 2))
            elements_to_insert.append(('para', ''))

            for level, suffix, title_es, title_en in sections:
                title = title_es if lang == "es" else title_en
                num = f"{base_num}{suffix}"
                elements_to_insert.append(('heading', f"{num} {title}", min(level + 1, 3)))
                elements_to_insert.append(('para', ''))

            # Create elements and insert in reverse order before ref_element
            for tag, text, *args in reversed(elements_to_insert):
                if tag == 'heading':
                    h_level = args[0] if args else 2
                    p = doc.add_heading(text, level=h_level)
                    # Move it before the reference element
                    ref_element.addprevious(p._element)
                else:
                    p = doc.add_paragraph(text)
                    ref_element.addprevious(p._element)
        else:
            # Append at the end
            heading_text = f"{base_num} Sprint {sprint_number}"
            doc.add_heading(heading_text, level=2)
            doc.add_paragraph("")

            for level, suffix, title_es, title_en in sections:
                title = title_es if lang == "es" else title_en
                num = f"{base_num}{suffix}"
                doc.add_heading(f"{num} {title}", level=min(level + 1, 3))
                doc.add_paragraph("")

    # ── Heading finding ──────────────────────────────────────────────

    def _find_heading(self, doc, search_text: str) -> Tuple[Optional[int], Optional[int]]:
        """
        Find a heading paragraph by partial text match.
        Returns (paragraph_index, heading_level) or (None, None).
        """
        search_norm = self._normalise_heading(search_text)
        for i, p in enumerate(doc.paragraphs):
            if p.style.name.startswith('Heading'):
                if search_norm in self._normalise_heading(p.text):
                    return i, self._heading_level(p)
        return None, None

    def _heading_level(self, para) -> int:
        """Return the numeric heading level (1, 2, 3...) of a paragraph."""
        style_name = para.style.name  # e.g. 'Heading 2'
        match = re.search(r'\d+', style_name)
        return int(match.group()) if match else 9

    def _normalise_heading(self, text: str) -> str:
        """Lowercase, strip whitespace and punctuation for heading comparison."""
        return re.sub(r'[^\w\s]', '', text.lower()).strip()

    def _clear_section_body(self, doc, heading_idx: int, heading_level: int) -> None:
        """Remove all paragraphs between heading_idx and the next heading of equal/higher level."""
        to_remove = []
        for j in range(heading_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            if p.style.name.startswith('Heading') and self._heading_level(p) <= heading_level:
                break
            to_remove.append(p)
        for p in to_remove:
            elem = p._element
            elem.getparent().remove(elem)

    # ── Paragraph insertion ──────────────────────────────────────────

    def _insert_paragraph_after(self, ref_para, text: str, heading_level: int = 0):
        """Insert a new paragraph right after *ref_para* and return the new Paragraph."""
        new_el = parse_xml(
            f'<w:p {nsdecls("w")}><w:r><w:t xml:space="preserve">{self._xml_escape(text)}</w:t></w:r></w:p>'
        )
        ref_para._element.addnext(new_el)
        # Wrap in a proper Paragraph object
        from docx.text.paragraph import Paragraph
        new_p = Paragraph(new_el, ref_para._element.getparent())
        if heading_level:
            new_p.style = ref_para._element.getparent().part.document.styles[f'Heading {heading_level}']
        return new_p

    def _xml_escape(self, text: str) -> str:
        """Escape XML special characters."""
        return (
            text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;")
        )

    # ── Table insertion ──────────────────────────────────────────────

    def _insert_table_after(self, doc, ref_para, headers: List[str], rows: List[List[str]]):
        """Insert a formatted table right after *ref_para*."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Headers
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
            # Header background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
            cell._element.get_or_add_tcPr().append(shading)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            for c_idx, value in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(value)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

        # Move the table element right after ref_para in the XML tree
        ref_para._element.addnext(table._tbl)

        return table

    # ── User Story card ──────────────────────────────────────────────

    def _insert_hu_card(self, doc, ref_para, story: Dict, lang: str):
        """Insert a formatted HU card (as a 2-column table) after ref_para."""
        sid = story.get("id", "HU??")
        name = story.get("name", "")
        priority = story.get("priority", "Media")
        sp = story.get("story_points", "")
        status = story.get("status", "Pendiente" if lang == "es" else "Pending")
        as_role = story.get("as_role", "")
        i_want = story.get("i_want", "")
        so_that = story.get("so_that", "")
        criteria = story.get("acceptance_criteria", [])
        developer = story.get("developer", "")

        rows_data = []
        if lang == "es":
            rows_data.append(("ID", sid))
            rows_data.append(("Nombre", name))
            rows_data.append(("Prioridad", priority))
            rows_data.append(("Story Points (PHU)", str(sp)))
            rows_data.append(("Estado", status))
            rows_data.append(("Como", as_role))
            rows_data.append(("Quiero", i_want))
            rows_data.append(("Para", so_that))
            rows_data.append(("Criterios de Aceptación", "\n".join(f"• {c}" for c in criteria) if criteria else ""))
            rows_data.append(("Desarrollador", developer))
        else:
            rows_data.append(("ID", sid))
            rows_data.append(("Name", name))
            rows_data.append(("Priority", priority))
            rows_data.append(("Story Points", str(sp)))
            rows_data.append(("Status", status))
            rows_data.append(("As a", as_role))
            rows_data.append(("I want", i_want))
            rows_data.append(("So that", so_that))
            rows_data.append(("Acceptance Criteria", "\n".join(f"• {c}" for c in criteria) if criteria else ""))
            rows_data.append(("Developer", developer))

        table = doc.add_table(rows=len(rows_data), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, (label, value) in enumerate(rows_data):
            # Label cell
            cell_l = table.rows[r_idx].cells[0]
            cell_l.text = label
            for run in cell_l.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E2EFDA" w:val="clear"/>')
            cell_l._element.get_or_add_tcPr().append(shading)

            # Value cell
            cell_v = table.rows[r_idx].cells[1]
            cell_v.text = str(value)
            for run in cell_v.paragraphs[0].runs:
                run.font.size = Pt(10)

        # Move table after ref_para
        ref_para._element.addnext(table._tbl)

        # Add spacing paragraph after the card
        spacer = doc.add_paragraph("")
        table._tbl.addnext(spacer._element)

        return spacer
